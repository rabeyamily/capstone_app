"""
DOCX parsing service using python-docx and docx2txt.
"""
import io
from typing import Tuple, Optional
from docx import Document
import docx2txt
from app.utils.text_cleaning import clean_text, normalize_whitespace, remove_encoding_issues


class DOCXParser:
    """DOCX parsing service using python-docx and docx2txt."""
    
    @staticmethod
    def extract_text(docx_content: bytes) -> Tuple[str, Optional[str]]:
        """
        Extract text from DOCX content.
        
        Args:
            docx_content: DOCX file content as bytes
            
        Returns:
            Tuple of (extracted_text, error_message)
        """
        try:
            # Create file-like object from bytes
            docx_file = io.BytesIO(docx_content)
            
            # Try using python-docx first (better structure preservation)
            try:
                doc = Document(docx_file)
                full_text = []
                
                # Extract text from all paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(paragraph.text.strip())
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            full_text.append(" | ".join(row_text))
                
                combined_text = "\n\n".join(full_text)
                
            except Exception as e:
                # Fallback to docx2txt if python-docx fails
                try:
                    docx_file.seek(0)  # Reset file pointer
                    combined_text = docx2txt.process(docx_file)
                except Exception as e2:
                    error_message = f"Error parsing DOCX with both methods: python-docx ({str(e)}), docx2txt ({str(e2)})"
                    return "", error_message
            
            if not combined_text or combined_text.strip() == "":
                return "", "No text could be extracted from the DOCX file"
            
            # Clean and normalize text
            cleaned_text = remove_encoding_issues(combined_text)
            cleaned_text = normalize_whitespace(cleaned_text)
            cleaned_text = clean_text(cleaned_text)
            
            return cleaned_text, None
            
        except Exception as e:
            error_message = f"Error parsing DOCX: {str(e)}"
            return "", error_message
    
    @staticmethod
    def extract_text_with_structure(docx_content: bytes) -> Tuple[str, Optional[str], dict]:
        """
        Extract text from DOCX with structure information.
        
        Args:
            docx_content: DOCX file content as bytes
            
        Returns:
            Tuple of (extracted_text, error_message, metadata)
        """
        try:
            docx_file = io.BytesIO(docx_content)
            metadata = {
                "total_paragraphs": 0,
                "total_tables": 0,
                "paragraphs_with_text": 0,
                "extraction_method": "python-docx"
            }
            
            try:
                doc = Document(docx_file)
                
                full_text = []
                paragraphs = []
                
                # Extract paragraphs
                for paragraph in doc.paragraphs:
                    metadata["total_paragraphs"] += 1
                    if paragraph.text.strip():
                        paragraphs.append(paragraph.text.strip())
                        metadata["paragraphs_with_text"] += 1
                
                full_text.extend(paragraphs)
                
                # Extract tables
                metadata["total_tables"] = len(doc.tables)
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            full_text.append(" | ".join(row_text))
                
                combined_text = "\n\n".join(full_text)
                
            except Exception as e:
                # Fallback to docx2txt
                try:
                    docx_file.seek(0)
                    combined_text = docx2txt.process(docx_file)
                    metadata["extraction_method"] = "docx2txt"
                    metadata["total_paragraphs"] = len(combined_text.split('\n'))
                except Exception as e2:
                    error_message = f"Error parsing DOCX: {str(e)}"
                    return "", error_message, metadata
            
            if not combined_text or combined_text.strip() == "":
                return "", "No text could be extracted from the DOCX file", metadata
            
            # Clean text
            cleaned_text = remove_encoding_issues(combined_text)
            cleaned_text = normalize_whitespace(cleaned_text)
            cleaned_text = clean_text(cleaned_text)
            
            return cleaned_text, None, metadata
            
        except Exception as e:
            error_message = f"Error parsing DOCX: {str(e)}"
            return "", error_message, {}
    
    @staticmethod
    def get_docx_metadata(docx_content: bytes) -> dict:
        """
        Extract metadata from DOCX.
        
        Args:
            docx_content: DOCX file content as bytes
            
        Returns:
            Dictionary with DOCX metadata
        """
        try:
            docx_file = io.BytesIO(docx_content)
            doc = Document(docx_file)
            
            metadata = {
                "total_paragraphs": len(doc.paragraphs),
                "total_tables": len(doc.tables),
                "core_properties": {
                    "title": doc.core_properties.title or "",
                    "author": doc.core_properties.author or "",
                    "subject": doc.core_properties.subject or "",
                    "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                    "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "",
                }
            }
            
            return metadata
            
        except Exception as e:
            return {
                "error": str(e),
                "total_paragraphs": 0,
                "total_tables": 0,
                "core_properties": {}
            }


# Global DOCX parser instance
docx_parser = DOCXParser()

